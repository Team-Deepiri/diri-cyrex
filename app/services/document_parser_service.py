"""
Document Parser Service for Spreadsheet Import
Parses various document types and extracts structured data for spreadsheet population
"""
from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path
import tempfile
import os
import io
from enum import Enum
import logging

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    """Supported document types"""
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    CSV = "csv"
    TXT = "txt"
    MD = "md"
    PNG = "png"
    JPG = "jpg"
    JPEG = "jpeg"
    TIFF = "tiff"
    HTML = "html"
    UNKNOWN = "unknown"


class ParsedDocument:
    """Structured representation of parsed document"""
    def __init__(self):
        self.document_type: Optional[DocumentType] = None
        self.tables: List[List[List[str]]] = []  # List of tables, each table is 2D array
        self.text_sections: List[Dict[str, Any]] = []  # Structured text sections
        self.key_value_pairs: Dict[str, str] = {}  # Extracted key-value pairs
        self.raw_text: str = ""  # Full extracted text
        self.metadata: Dict[str, Any] = {}  # Document metadata
        self.confidence_scores: Dict[str, float] = {}  # Confidence scores for extraction
        self.document_category: Optional[str] = None  # ML-classified category (invoice, receipt, etc.)
        self.layout_elements: List[Dict[str, Any]] = []  # Layout structure
        self.column_mapping: Optional[Dict[str, Any]] = None  # Smart column mapping suggestions


class DocumentParserService:
    """Service for parsing documents into spreadsheet-ready format"""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.xlsx': DocumentType.XLSX,
            '.csv': DocumentType.CSV,
            '.txt': DocumentType.TXT,
            '.md': DocumentType.MD,
            '.png': DocumentType.PNG,
            '.jpg': DocumentType.JPG,
            '.jpeg': DocumentType.JPEG,
            '.tiff': DocumentType.TIFF,
            '.html': DocumentType.HTML,
        }
    
    def detect_document_type(self, filename: str, content: Optional[bytes] = None) -> DocumentType:
        """Detect document type from filename or content"""
        ext = Path(filename).suffix.lower()
        doc_type = self.supported_types.get(ext, DocumentType.UNKNOWN)
        
        # If unknown, try to detect from content
        if doc_type == DocumentType.UNKNOWN and content:
            # Check for PDF magic bytes
            if content[:4] == b'%PDF':
                return DocumentType.PDF
            # Check for ZIP (DOCX/XLSX are ZIP files)
            if content[:2] == b'PK':
                # Could be DOCX or XLSX, check filename
                if ext in ['.docx', '.xlsx']:
                    return self.supported_types.get(ext, DocumentType.UNKNOWN)
        
        return doc_type
    
    async def parse_document(
        self,
        file_content: bytes,
        filename: str,
        use_ocr: bool = True,
        extract_tables: bool = True,
        detect_layout: bool = False,
        ocr_language: Optional[str] = None,
        use_doclaynet: bool = False,
    ) -> ParsedDocument:
        """
        Parse a document and extract structured data
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            use_ocr: Whether to use OCR for images/scanned PDFs
            extract_tables: Whether to extract tables
            detect_layout: Whether to perform layout analysis
        
        Returns:
            ParsedDocument with extracted data
        """
        parsed = ParsedDocument()
        doc_type = self.detect_document_type(filename, file_content)
        parsed.document_type = doc_type
        parsed.metadata['filename'] = filename
        parsed.metadata['detected_type'] = doc_type.value
        
        logger.info(f"Parsing document: {filename}, type: {doc_type.value}")
        
        # Save to temp file for parsing
        suffix = Path(filename).suffix or '.tmp'
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(file_content)
            temp_path = temp_file.name
        
        try:
            if doc_type == DocumentType.PDF:
                await self._parse_pdf(temp_path, parsed, use_ocr, extract_tables, ocr_language, use_doclaynet)
            elif doc_type == DocumentType.DOCX:
                await self._parse_docx(temp_path, parsed, extract_tables)
            elif doc_type in [DocumentType.XLSX, DocumentType.CSV]:
                await self._parse_spreadsheet(temp_path, parsed, doc_type)
            elif doc_type in [DocumentType.TXT, DocumentType.MD]:
                await self._parse_text(temp_path, parsed)
            elif doc_type in [DocumentType.PNG, DocumentType.JPG, DocumentType.JPEG, DocumentType.TIFF]:
                await self._parse_image(temp_path, parsed, use_ocr, ocr_language)
            elif doc_type == DocumentType.HTML:
                await self._parse_html(temp_path, parsed)
            else:
                raise ValueError(f"Unsupported document type: {doc_type.value}")
            
            # Advanced analysis if enabled
            if detect_layout or parsed.raw_text:
                try:
                    from .advanced_document_analyzer import get_advanced_analyzer, DocumentCategory
                    analyzer = get_advanced_analyzer()
                    
                    # Classify document
                    category, category_confidence = await analyzer.classify_document(
                        parsed.raw_text,
                        filename,
                        use_llm=True,
                    )
                    parsed.document_category = category.value
                    parsed.metadata['category'] = category.value
                    parsed.metadata['category_confidence'] = category_confidence
                    
                    # Extract key-value pairs
                    if parsed.raw_text:
                        kvp = await analyzer.extract_key_value_pairs(
                            parsed.raw_text,
                            category,
                            use_llm=True,
                        )
                        parsed.key_value_pairs.update(kvp)
                    
                    # Detect layout
                    if detect_layout:
                        layout_elements = await analyzer.detect_layout(
                            parsed.raw_text,
                            parsed.tables,
                            use_ml=False,  # Future: enable ML-based layout detection
                        )
                        parsed.layout_elements = [
                            {
                                'type': el.element_type,
                                'content': el.content[:200],  # Truncate for response
                                'confidence': el.confidence,
                            }
                            for el in layout_elements[:50]  # Limit to 50 elements
                        ]
                    
                    # Get column mapping suggestions
                    parsed.column_mapping = analyzer.get_column_mapping(
                        category,
                        parsed.tables,
                        parsed.key_value_pairs,
                    )
                    
                except Exception as e:
                    logger.warning(f"Advanced analysis failed: {e}", exc_info=True)
                    # Continue without advanced features
            
            # Set confidence scores
            parsed.confidence_scores = {
                'text_extraction': 0.9 if parsed.raw_text else 0.0,
                'table_extraction': 0.8 if parsed.tables else 0.0,
                'classification': parsed.metadata.get('category_confidence', 0.0),
                'overall': 0.85 if (parsed.raw_text or parsed.tables) else 0.0,
            }
            
        finally:
            # Clean up temp file
            if os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temp file: {e}")
        
        return parsed
    
    async def _parse_pdf(
        self,
        file_path: str,
        parsed: ParsedDocument,
        use_ocr: bool,
        extract_tables: bool,
        ocr_language: Optional[str] = None,
        use_doclaynet: bool = False,
    ):
        """Parse PDF document with optional DocLayNet layout detection"""
        try:
            # Try DocLayNet for advanced layout detection if enabled
            if use_doclaynet:
                try:
                    layout_result = await self._parse_with_doclaynet(file_path)
                    if layout_result:
                        parsed.raw_text = layout_result.get('text', '')
                        parsed.tables = layout_result.get('tables', [])
                        parsed.layout_elements = layout_result.get('layout_elements', [])
                        parsed.metadata['doclaynet_used'] = True
                        parsed.metadata['page_count'] = layout_result.get('page_count', 0)
                        return
                except Exception as e:
                    logger.warning(f"DocLayNet parsing failed, falling back to standard: {e}")
            
            # Standard PDF parsing
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                # Extract text from all pages
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                parsed.raw_text = '\n\n'.join(text_parts)
                
                # Extract tables if requested
                if extract_tables:
                    for page in pdf.pages:
                        tables = page.extract_tables()
                        for table in tables:
                            if table and len(table) > 0:
                                # Clean up table data
                                cleaned_table = []
                                for row in table:
                                    if row:
                                        cleaned_row = [str(cell) if cell is not None else '' for cell in row]
                                        cleaned_table.append(cleaned_row)
                                if cleaned_table:
                                    parsed.tables.append(cleaned_table)
                
                parsed.metadata['page_count'] = len(pdf.pages)
                
        except ImportError:
            logger.error("pdfplumber not installed. Install with: pip install pdfplumber")
            raise
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}", exc_info=True)
            # Try OCR fallback if text extraction failed
            if use_ocr and not parsed.raw_text:
                await self._parse_pdf_with_ocr(file_path, parsed, ocr_language)
    
    async def _parse_with_doclaynet(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Parse PDF using DocLayNet for advanced layout detection"""
        try:
            # DocLayNet integration (requires docling library)
            # This is a placeholder - actual implementation would use docling library
            # For now, we'll prepare the structure for future integration
            
            # Future implementation would look like:
            # from docling import DocumentConverter
            # converter = DocumentConverter()
            # doc = converter.convert(file_path)
            # return {
            #     'text': doc.text,
            #     'tables': [table.to_dict() for table in doc.tables],
            #     'layout_elements': [el.to_dict() for el in doc.layout_elements],
            #     'page_count': len(doc.pages),
            # }
            
            logger.info("DocLayNet integration prepared but not yet implemented. Install docling library for advanced layout detection.")
            return None
            
        except Exception as e:
            logger.warning(f"DocLayNet parsing not available: {e}")
            return None
    
    async def _parse_pdf_with_ocr(self, file_path: str, parsed: ParsedDocument, language: Optional[str] = None):
        """Parse PDF using OCR (for scanned PDFs) with multi-language support"""
        try:
            from pdf2image import convert_from_path
            import pytesseract
            from PIL import Image
            
            # Convert PDF to images
            images = convert_from_path(file_path)
            text_parts = []
            
            # Detect language from first page if not provided
            if not language and images:
                language = await self._detect_ocr_language(images[0])
            
            for img in images:
                # Perform OCR with language
                if language and language != 'eng':
                    text = pytesseract.image_to_string(img, lang=language)
                else:
                    text = pytesseract.image_to_string(img, lang='eng')
                if text:
                    text_parts.append(text)
            
            parsed.raw_text = '\n\n'.join(text_parts)
            parsed.metadata['ocr_used'] = True
            parsed.metadata['ocr_language'] = language or 'eng'
            parsed.confidence_scores['ocr'] = 0.7  # OCR is less accurate
            
        except ImportError:
            logger.warning("OCR dependencies not installed. Install with: pip install pdf2image pytesseract Pillow")
        except Exception as e:
            logger.error(f"Error in OCR: {e}", exc_info=True)
    
    async def _parse_docx(self, file_path: str, parsed: ParsedDocument, extract_tables: bool):
        """Parse Word document"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            parsed.raw_text = '\n'.join(paragraphs)
            
            # Extract tables if requested
            if extract_tables:
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        if any(row_data):  # Skip empty rows
                            table_data.append(row_data)
                    if table_data:
                        parsed.tables.append(table_data)
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}", exc_info=True)
            raise
    
    async def _parse_spreadsheet(
        self,
        file_path: str,
        parsed: ParsedDocument,
        doc_type: DocumentType,
    ):
        """Parse Excel or CSV file"""
        try:
            import pandas as pd
            
            if doc_type == DocumentType.CSV:
                # Try different encodings
                encodings = ['utf-8', 'latin-1', 'cp1252']
                df = None
                for encoding in encodings:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                
                if df is None:
                    raise ValueError("Could not decode CSV file")
            else:  # XLSX
                df = pd.read_excel(file_path, engine='openpyxl')
            
            # Convert DataFrame to table format
            table_data = []
            # Add header row
            table_data.append(df.columns.tolist())
            # Add data rows
            for _, row in df.iterrows():
                table_data.append([str(val) if pd.notna(val) else '' for val in row.tolist()])
            
            parsed.tables.append(table_data)
            parsed.raw_text = df.to_string()
            parsed.metadata['row_count'] = len(df)
            parsed.metadata['column_count'] = len(df.columns)
            
        except ImportError:
            logger.error("pandas/openpyxl not installed. Install with: pip install pandas openpyxl")
            raise
        except Exception as e:
            logger.error(f"Error parsing spreadsheet: {e}", exc_info=True)
            raise
    
    async def _parse_text(self, file_path: str, parsed: ParsedDocument):
        """Parse plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError("Could not decode text file")
            
            parsed.raw_text = text
            
        except Exception as e:
            logger.error(f"Error parsing text file: {e}", exc_info=True)
            raise
    
    async def _parse_image(
        self,
        file_path: str,
        parsed: ParsedDocument,
        use_ocr: bool,
        language: Optional[str] = None,
    ):
        """Parse image file using OCR with multi-language support"""
        if not use_ocr:
            raise ValueError("OCR required for image files")
        
        try:
            import pytesseract
            from PIL import Image
            
            img = Image.open(file_path)
            
            # Detect language if not provided
            if not language:
                language = await self._detect_ocr_language(img)
            
            # Perform OCR with language
            if language and language != 'eng':
                # Multi-language OCR (e.g., 'eng+spa' for English + Spanish)
                text = pytesseract.image_to_string(img, lang=language)
            else:
                # Default English OCR
                text = pytesseract.image_to_string(img, lang='eng')
            
            parsed.raw_text = text
            parsed.metadata['ocr_used'] = True
            parsed.metadata['ocr_language'] = language or 'eng'
            parsed.confidence_scores['ocr'] = 0.7
            
        except ImportError:
            logger.error("pytesseract/Pillow not installed. Install with: pip install pytesseract Pillow")
            raise
        except Exception as e:
            logger.error(f"Error parsing image with OCR: {e}", exc_info=True)
            raise
    
    async def _detect_ocr_language(self, image) -> str:
        """Detect language for OCR (simple heuristic, can be enhanced with ML)"""
        try:
            import pytesseract
            
            # Try to detect language using Tesseract's language detection
            # This is a simple approach - can be enhanced with proper language detection models
            # For now, try common languages
            common_languages = ['eng', 'spa', 'fra', 'deu', 'por', 'chi_sim', 'jpn', 'kor']
            
            # Quick test with first 100 characters
            test_text = pytesseract.image_to_string(image, lang='+'.join(common_languages))
            
            # Simple heuristic: if we get good text, use multi-language
            # Otherwise default to English
            if len(test_text.strip()) > 10:
                # Use multi-language OCR
                return '+'.join(common_languages[:3])  # Limit to 3 languages for performance
            else:
                return 'eng'
                
        except Exception as e:
            logger.warning(f"Language detection failed: {e}, defaulting to English")
            return 'eng'
    
    async def _parse_html(self, file_path: str, parsed: ParsedDocument):
        """Parse HTML file"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract text
            parsed.raw_text = soup.get_text(separator='\n', strip=True)
            
            # Extract tables
            tables = soup.find_all('table')
            for table in tables:
                table_data = []
                rows = table.find_all('tr')
                for row in rows:
                    cells = row.find_all(['td', 'th'])
                    row_data = [cell.get_text(strip=True) for cell in cells]
                    if any(row_data):
                        table_data.append(row_data)
                if table_data:
                    parsed.tables.append(table_data)
            
        except ImportError:
            logger.warning("BeautifulSoup not installed. Install with: pip install beautifulsoup4")
            # Fallback to basic text extraction
            with open(file_path, 'r', encoding='utf-8') as f:
                parsed.raw_text = f.read()
        except Exception as e:
            logger.error(f"Error parsing HTML: {e}", exc_info=True)
    
    def to_spreadsheet_format(
        self,
        parsed: ParsedDocument,
        start_cell: str = "A1",
    ) -> Dict[str, Any]:
        """
        Convert parsed document to spreadsheet format
        
        Args:
            parsed: ParsedDocument instance
            start_cell: Starting cell (e.g., "A1")
        
        Returns:
            Dictionary with spreadsheet data mapping
        """
        # Determine best format based on content
        if parsed.tables:
            # Use first table (or combine all tables)
            if len(parsed.tables) == 1:
                data = parsed.tables[0]
            else:
                # Combine multiple tables with separator rows
                data = []
                for i, table in enumerate(parsed.tables):
                    if i > 0:
                        # Add separator row
                        data.append(['---', '---', '---'])
                    data.extend(table)
        else:
            # Convert text to rows
            lines = parsed.raw_text.split('\n')
            data = [[line] for line in lines if line.strip()]
        
        return {
            "start_cell": start_cell,
            "data": data,
            "has_tables": len(parsed.tables) > 0,
            "table_count": len(parsed.tables),
            "row_count": len(data),
            "column_count": max(len(row) for row in data) if data else 0,
        }


# Singleton instance
_parser_service: Optional[DocumentParserService] = None


def get_document_parser_service() -> DocumentParserService:
    """Get singleton DocumentParserService instance"""
    global _parser_service
    if _parser_service is None:
        _parser_service = DocumentParserService()
    return _parser_service

