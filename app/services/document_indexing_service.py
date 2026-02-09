"""
Document Indexing Service
Handles document ingestion, processing, chunking, and indexing into Milvus
for the Language Intelligence Platform
"""

from typing import List, Dict, Optional, Any
from enum import Enum
from pathlib import Path
from dataclasses import dataclass
from datetime import datetime

# LangChain imports
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Local imports
from ..logging_config import get_logger
from ..integrations.milvus_store import get_milvus_store
from .document_processors import LeaseProcessor, ContractProcessor

logger = get_logger("cyrex.document_indexing")

# Document format enum - matches MD file (must be defined before IndexingResult)
class DocumentFormat(str, Enum):
    """Supported file formats"""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"
    CSV = "csv"

# Document type enum 
class B2BDocumentType(str, Enum):
    """Document types for Language Intelligence Platform"""
    LEGAL_DOCUMENT = "legal_document"  # General legal document
    CONTRACT = "contract"               # Contract documents
    LEASE = "lease"                     # Lease documents
    REGULATION = "regulation"           # Regulatory documents
    AMENDMENT = "amendment"             # Amendments
    COMPLIANCE_REPORT = "compliance_report"  # Compliance reports
    TEMPLATE = "template"               # Templates (for drift detection)

@dataclass
class IndexingResult:
    """Result object returned from indexing operations"""
    document_id: str
    title: str
    chunk_count: int
    format: DocumentFormat
    success: bool = True
    file_size: Optional[int] = None
    indexed_at: Optional[datetime] = None
    
    def __post_init__(self):
        if self.indexed_at is None:
            self.indexed_at = datetime.now()

class DocumentIndexingService:
    """
    Service for indexing documents into Milvus for RAG retrieval.
    Handles file parsing, chunking, embedding, and indexing.
    Supports the Language Intelligence Platform use cases.
    """
    
    def __init__(
        self,
        collection_name: str, # Required: Name of the Milvus collection
        chunk_size: int = 1500, # Optional: Size of text chunks (default 1500 for legal documents)
        chunk_overlap: int = 300, # Optional: Overlap between chunks (default 300 for clause preservation)
        chunking_strategy: str = "paragraph", # Optional: Chunking strategy - "paragraph" (preserves clauses), "sentence", or "character"
        use_enhanced_rag: bool = True, # Optional: Whether to use enhanced RAG features
        enable_auto_extraction: bool = True # Optional: Whether to automatically extract obligations and clauses
    ):
        """
        Initialize DocumentIndexingService.
        
        Args:
            collection_name: Name of the Milvus collection
            chunk_size: Size of text chunks (default 1500 for legal documents)
            chunk_overlap: Overlap between chunks (default 300 for clause preservation)
            chunking_strategy: Chunking strategy - "paragraph" (preserves clauses), 
                             "sentence", or "character"
            use_enhanced_rag: Whether to use enhanced RAG features
        """
        self.collection_name = collection_name
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.chunking_strategy = chunking_strategy
        self.use_enhanced_rag = use_enhanced_rag
        self.enable_auto_extraction = enable_auto_extraction
        
        # Initialize processors (lazy-loaded when needed)
        self._lease_processor = None
        self._contract_processor = None
        
        # Initialize vector store (connects to Milvus)
        try:
            self.vector_store = get_milvus_store(
                collection_name=collection_name
            )
            logger.info(
                "Vector store initialized",
                collection=collection_name
            )
        except Exception as e:
            logger.error(
                "Failed to initialize vector store",
                collection=collection_name,
                error=str(e)
            )
            raise
        
        # Initialize text splitter based on strategy
        # Paragraph strategy preserves clause boundaries (important for legal docs)
        if chunking_strategy == "paragraph":
            separators = ["\n\n", "\n", " ", ""]
        elif chunking_strategy == "sentence":
            separators = ["\n", ". ", " ", ""]
        else:  # character
            separators = [" ", ""]
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators,
            length_function=len
        )
        
        logger.info(
            "DocumentIndexingService initialized",
            collection=collection_name,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            strategy=chunking_strategy
        )

    
    def _parse_file(self, file_path: str) -> str:
        """
        Parse file and extract text content.
        
        Args:
            file_path: Path to the file
        
        Returns:
            Extracted text content
        """
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path_obj.suffix.lower().lstrip('.')
        logger.info("Parsing file", file_path=file_path, format=file_ext)
        
        try:
            if file_ext == 'pdf':
                return self._parse_pdf(file_path)
            elif file_ext in ['docx', 'doc']:
                return self._parse_docx(file_path)
            elif file_ext in ['txt', 'md']:
                return self._parse_txt(file_path)
            elif file_ext == 'csv':
                return self._parse_csv(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
        except Exception as e:
            logger.error("Failed to parse file", file_path=file_path, error=str(e))
            raise

    def _parse_csv(self, file_path: str) -> str:
        """Parse CSV file"""
        try:
            import csv
            text_lines = []
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    text_lines.append(" | ".join(row))  # Join columns with pipe separator
            text = "\n".join(text_lines)
            logger.info("Extracted text from CSV", file_path=file_path, length=len(text))
            return text
        except Exception as e:
            logger.error("Failed to parse CSV", file_path=file_path, error=str(e))
            raise
    
    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file"""
        import sys
        import os
        import importlib.util
        
        # Ensure site-packages is in the path (add it first for priority)
        site_packages = '/opt/conda/lib/python3.11/site-packages'
        if site_packages not in sys.path:
            sys.path.insert(0, site_packages)
            logger.debug("Added site-packages to sys.path", site_packages=site_packages)
        
        # Try multiple import strategies
        pdfplumber = None
        import_err = None
        
        # Strategy 1: Try direct importlib
        try:
            import importlib
            if 'pdfplumber' in sys.modules:
                del sys.modules['pdfplumber']
            pdfplumber = importlib.import_module('pdfplumber')
            logger.debug("pdfplumber imported via importlib")
        except Exception as e:
            import_err = e
            logger.debug("importlib import failed, trying spec loader", error=str(e))
            
            # Strategy 2: Try using spec loader (more explicit)
            try:
                spec_path = os.path.join(site_packages, 'pdfplumber', '__init__.py')
                if os.path.exists(spec_path):
                    spec = importlib.util.spec_from_file_location("pdfplumber", spec_path)
                    if spec and spec.loader:
                        pdfplumber = importlib.util.module_from_spec(spec)
                        spec.loader.exec_module(pdfplumber)
                        logger.debug("pdfplumber imported via spec loader")
                    else:
                        raise ImportError("Could not create spec for pdfplumber")
                else:
                    raise ImportError(f"pdfplumber __init__.py not found at {spec_path}")
            except Exception as e2:
                import_err = e2
                logger.debug("spec loader import failed", error=str(e2))
        
        # If all strategies failed, log and raise
        if pdfplumber is None:
            logger.error(
                "Failed to import pdfplumber",
                error_type=type(import_err).__name__ if import_err else 'Unknown',
                error_message=str(import_err) if import_err else 'All import strategies failed',
                python_executable=sys.executable,
                python_version=sys.version,
                python_path=sys.path,
                pythonpath_env=os.environ.get('PYTHONPATH', 'not set'),
                site_packages_exists=os.path.exists(site_packages),
                pdfplumber_dir_exists=os.path.exists(os.path.join(site_packages, 'pdfplumber')),
                exc_info=True
            )
            raise ImportError(
                f"pdfplumber not available. Error: {import_err}. "
                f"Python: {sys.executable}, Path: {sys.path}, PYTHONPATH: {os.environ.get('PYTHONPATH', 'not set')}"
            ) from import_err
        
        logger.debug("pdfplumber imported successfully", 
                   pdfplumber_version=getattr(pdfplumber, '__version__', 'unknown'),
                   python_executable=sys.executable)
        
        # Now try to parse the PDF
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:  # Only add non-None, non-empty text
                        text_parts.append(page_text)
                text = "\n".join(text_parts)
            
            if not text or len(text.strip()) == 0:
                logger.warning("PDF extracted but contains no text", file_path=file_path)
                # Return empty string rather than raising error
                return ""
            
            logger.info("Extracted text from PDF", file_path=file_path, length=len(text))
            return text
        except Exception as e:
            # Check if this is actually an ImportError from a dependency
            error_type = type(e).__name__
            error_msg = str(e)
            
            # Check if the error mentions pdfplumber or a missing module
            if "pdfplumber" in error_msg.lower() or "no module named" in error_msg.lower():
                logger.error(
                    "PDF parsing failed - possible import issue",
                    error_type=error_type,
                    error_message=error_msg,
                    file_path=file_path,
                    python_executable=sys.executable,
                    python_path=sys.path,
                    pythonpath_env=os.environ.get('PYTHONPATH', 'not set'),
                    exc_info=True
                )
                raise ImportError(
                    f"pdfplumber or its dependency not available: {error_msg}. "
                    f"Python: {sys.executable}, Path: {sys.path}, PYTHONPATH: {os.environ.get('PYTHONPATH', 'not set')}"
                ) from e
            
            logger.error("Failed to parse PDF", file_path=file_path, error=str(e), exc_info=True)
            raise
    
    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file"""
        try:
            import docx
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            logger.info("Extracted text from DOCX", file_path=file_path, length=len(text))
            return text
        except Exception as e:
            logger.error("Failed to parse DOCX", file_path=file_path, error=str(e))
            raise
    
    def _parse_txt(self, file_path: str) -> str:
        """Parse TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            logger.info("Extracted text from TXT", file_path=file_path, length=len(text))
            return text
        except Exception as e:
            logger.error("Failed to parse TXT", file_path=file_path, error=str(e))
            raise
    
    def _chunk_document(
        self,
        text: str,
        document_id: str,
        metadata: Dict[str, Any]
    ) -> List[Document]:
        """
        Split document text into smaller chunks for vector indexing.
        
        Takes the full document text and splits it into chunks of approximately
        chunk_size characters, preserving paragraph boundaries when possible.
        Each chunk is wrapped in a LangChain Document object with metadata.
        
        Args:
            text: The full document text to chunk
            document_id: Unique identifier for the document
            metadata: Document metadata to preserve in each chunk (version, 
                    document_type, etc.)
        
        Returns:
            List of Document objects, each containing:
            - page_content: A chunk of text (~chunk_size characters)
            - metadata: Original metadata plus chunk_index and total_chunks
        """
        documents = []
        chunks = self.text_splitter.split_text(text)
        for i, chunk in enumerate(chunks):
            chunk_metadata = {
                **metadata, # Business metadata from MD file (version, etc.)
                 # Technical metadata (not in MD file, but useful)
                "chunk_index": i,
                "document_id": document_id,
                "total_chunks": len(chunks)
            }
            documents.append(Document(page_content=chunk, metadata=chunk_metadata))
        return documents

    async def _generate_unique_document_id(self, title: str) -> str:
        """
        Generate a unique document_id from title, ensuring no duplicates.
        
        Args:
            title: Document title
            
        Returns:
            Unique document_id
        """
        import hashlib
        import uuid
        
        # Base ID from title hash
        base_id = f"doc_{hashlib.md5(title.encode()).hexdigest()[:12]}"
        
        # Check if this ID already exists
        existing_chunks = self.vector_store.list_documents(
            filters={"document_id": base_id},
            limit=1
        )
        
        if not existing_chunks:
            # ID is unique, return it
            return base_id
        
        # ID exists, append UUID suffix to ensure uniqueness
        unique_suffix = uuid.uuid4().hex[:8]
        unique_id = f"{base_id}_{unique_suffix}"
        
        logger.info(
            "Document ID collision detected, using unique suffix",
            base_id=base_id,
            unique_id=unique_id,
            title=title
        )
        
        return unique_id

    async def _extract_metadata_automatically(
        self,
        text: str,
        doc_type: B2BDocumentType,
        metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Automatically extract obligations and clauses from document text.
        
        Uses LLM-based processors to extract structured data from leases and contracts.
        Merges extracted data with manually provided metadata.
        
        Args:
            text: Document text content
            doc_type: Document type
            metadata: Existing metadata (manual inputs)
        
        Returns:
            Updated metadata with extracted obligations and clauses
        """
        if not self.enable_auto_extraction:
            return metadata
        
        extracted_metadata = {}
        
        try:
            # Extract based on document type
            if doc_type == B2BDocumentType.LEASE:
                if self._lease_processor is None:
                    from ..integrations.llm_providers import get_llm_provider
                    llm_provider = get_llm_provider()
                    self._lease_processor = LeaseProcessor(llm_provider=llm_provider)
                
                # Get lease_id from metadata for context
                lease_id = metadata.get("lease_id") or metadata.get("document_id", "")
                
                # Process lease document
                result = await self._lease_processor.process(
                    document_text=text,
                    document_url=f"file://{lease_id}",
                    lease_id=lease_id
                )
                
                # Map processor output to metadata schema
                obligations = result.get("obligations", [])
                clauses = result.get("keyClauses", [])
                
                # Convert obligations to metadata format
                if obligations:
                    extracted_metadata["obligations"] = [
                        {
                            "obligation_id": f"OBL-{i+1:03d}",
                            "type": ob.get("obligationType", "").lower() if ob.get("obligationType") else "other",
                            "deadline": ob.get("deadline", ""),
                            "owner": ob.get("party", "").lower() if ob.get("party") else "unknown",
                            "amount": ob.get("amount"),
                            "frequency": ob.get("frequency", "").lower() if ob.get("frequency") else None
                        }
                        for i, ob in enumerate(obligations)
                        if ob
                    ]
                
                # Convert clauses to metadata format
                if clauses:
                    extracted_metadata["clauses"] = [
                        {
                            "clause_id": f"CLAUSE-{i+1:03d}",
                            "type": clause.get("clauseType", "").lower() if clause.get("clauseType") else "other",
                            "section": clause.get("title", ""),
                            "text": clause.get("fullText", clause.get("summary", ""))
                        }
                        for i, clause in enumerate(clauses)
                        if clause
                    ]
                
                logger.info(
                    "Auto-extracted lease metadata",
                    obligations_count=len(extracted_metadata.get("obligations", [])),
                    clauses_count=len(extracted_metadata.get("clauses", []))
                )
            
            elif doc_type == B2BDocumentType.CONTRACT:
                if self._contract_processor is None:
                    from ..integrations.llm_providers import get_llm_provider
                    llm_provider = get_llm_provider()
                    self._contract_processor = ContractProcessor(llm_provider=llm_provider)
                
                # Get contract_id from metadata for context
                contract_id = metadata.get("contract_id") or metadata.get("document_id", "")
                
                # Process contract document
                result = await self._contract_processor.process(
                    document_text=text,
                    document_url=f"file://{contract_id}",
                    contract_id=contract_id
                )
                
                # Map processor output to metadata schema
                obligations = result.get("obligations", [])
                clauses = result.get("keyClauses", [])
                
                # Convert obligations to metadata format
                if obligations:
                    extracted_metadata["obligations"] = [
                        {
                            "obligation_id": f"OBL-{i+1:03d}",
                            "type": ob.get("obligationType", "").lower() if ob.get("obligationType") else "other",
                            "deadline": ob.get("deadline", ""),
                            "owner": ob.get("party", "").lower() if ob.get("party") else "unknown",
                            "amount": ob.get("amount"),
                            "depends_on": ob.get("dependencies", []) if ob.get("dependencies") else None
                        }
                        for i, ob in enumerate(obligations)
                        if ob
                    ]
                
                # Convert clauses to metadata format
                if clauses:
                    extracted_metadata["clauses"] = [
                        {
                            "clause_id": f"CLAUSE-{i+1:03d}",
                            "type": clause.get("clauseType", "").lower() if clause.get("clauseType") else "other",
                            "section": clause.get("section", clause.get("clauseNumber", "")),
                            "text": clause.get("clauseText", "")
                        }
                        for i, clause in enumerate(clauses)
                        if clause
                    ]
                
                logger.info(
                    "Auto-extracted contract metadata",
                    obligations_count=len(extracted_metadata.get("obligations", [])),
                    clauses_count=len(extracted_metadata.get("clauses", []))
                )
        
        except Exception as e:
            logger.warning(
                "Auto-extraction failed, using manual metadata only",
                doc_type=doc_type.value,
                error=str(e)
            )
            # Return original metadata if extraction fails
            return metadata
        
        # Merge extracted metadata with manual metadata
        # Manual metadata takes precedence if both exist
        merged_metadata = {**metadata}  # Start with manual metadata
        
        # Only add extracted obligations/clauses if not manually provided
        if extracted_metadata.get("obligations") and not metadata.get("obligations"):
            merged_metadata["obligations"] = extracted_metadata["obligations"]
            logger.info("Using auto-extracted obligations", count=len(extracted_metadata["obligations"]))
        elif extracted_metadata.get("obligations") and metadata.get("obligations"):
            logger.info("Manual obligations provided, skipping auto-extracted obligations")
        
        if extracted_metadata.get("clauses") and not metadata.get("clauses"):
            merged_metadata["clauses"] = extracted_metadata["clauses"]
            logger.info("Using auto-extracted clauses", count=len(extracted_metadata["clauses"]))
        elif extracted_metadata.get("clauses") and metadata.get("clauses"):
            logger.info("Manual clauses provided, skipping auto-extracted clauses")
        
        return merged_metadata

    async def index_file(
        self,
        file_path: str,
        document_id: Optional[str],
        title: str,
        doc_type: B2BDocumentType,
        industry: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Index a file into Milvus for RAG retrieval.
        
        Orchestrates the complete pipeline: parse → chunk → embed → store.
        Supports PDF, DOCX, TXT, MD, and CSV formats.
        
        Args:
            file_path: Path to the file to index
            document_id: Unique document identifier (will be generated if None)
            title: Document title
            doc_type: Document type (LEASE, CONTRACT, REGULATION, etc.)
            industry: Industry classification
            metadata: Optional metadata dict (version, document-specific fields, etc.)
                    See MD file for complete metadata schemas.
        
        Returns:
            Dict with document_id, title, chunk_count, format, and status.
        
        Raises:
                FileNotFoundError: If file doesn't exist
                ValueError: If file format is unsupported
                RuntimeError: If indexing fails
        """
        try:
            # Parse file
            text = self._parse_file(file_path)
            
            # Get file extension for return value
            file_ext = Path(file_path).suffix.lower().lstrip('.')
            
            # Generate document_id if not provided, ensuring uniqueness
            if not document_id or document_id is None:
                document_id = await self._generate_unique_document_id(title)
                logger.info(
                    "Generated document_id from title",
                    document_id=document_id,
                    title=title
                )
            
            # Build base metadata
            full_metadata = metadata or {}
            full_metadata.update({
                "document_id": document_id,  # Now always has a value
                "title": title,
                "doc_type": doc_type.value,
                "industry": industry
            })
            
            # Automatically extract obligations and clauses if enabled
            if self.enable_auto_extraction and doc_type in [B2BDocumentType.LEASE, B2BDocumentType.CONTRACT]:
                try:
                    full_metadata = await self._extract_metadata_automatically(
                        text=text,
                        doc_type=doc_type,
                        metadata=full_metadata
                    )
                except Exception as e:
                    logger.warning(
                        "Auto-extraction failed, continuing with manual metadata",
                        error=str(e)
                    )
                    # Continue with manual metadata only
            
            # Chunk document
            documents = self._chunk_document(text, document_id, full_metadata)
            
            # Store in Milvus
            self.vector_store.add_documents(documents)
            
            logger.info(
                "Document indexed successfully",
                document_id=document_id,
                chunk_count=len(documents),
                format=file_ext
            )
            
            # Get file size
            file_path_obj = Path(file_path)
            file_size = file_path_obj.stat().st_size if file_path_obj.exists() else None
            
            # Convert format string to enum (default to TXT if unknown)
            try:
                format_enum = DocumentFormat(file_ext)
            except ValueError:
                format_enum = DocumentFormat.TXT
                logger.warning(f"Unknown format '{file_ext}', defaulting to TXT")
            
            return IndexingResult(
                document_id=document_id,
                title=title,
                chunk_count=len(documents),
                format=format_enum,
                file_size=file_size
            )
            
        except FileNotFoundError as e:
            logger.error(
                "File not found",
                file_path=file_path,
                document_id=document_id,
                error=str(e)
            )
            raise  # Re-raise specific error
            
        except ValueError as e:
            logger.error(
                "Invalid file format or parameter",
                file_path=file_path,
                document_id=document_id,
                error=str(e)
            )
            raise  # Re-raise specific error
            
        except Exception as e:
            logger.error(
                "Failed to index document",
                file_path=file_path,
                document_id=document_id,
                error=str(e),
                exc_info=True  # Include stack trace
            )
            raise RuntimeError(f"Failed to index document: {str(e)}") from e

    async def index_text(
        self,
        text: str,
        document_id: Optional[str],
        title: str,
        doc_type: B2BDocumentType,
        industry: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Index text content directly into Milvus for RAG retrieval.
        
        Similar to index_file but takes text directly instead of parsing a file.
        
        Args:
            text: Text content to index
            document_id: Unique document identifier (will be generated if None)
            title: Document title
            doc_type: Document type (LEASE, CONTRACT, REGULATION, etc.)
            industry: Industry classification
            metadata: Optional metadata dict (version, document-specific fields, etc.)
        
        Returns:
            IndexingResult with document_id, title, chunk_count, and status.
        
        Raises:
            RuntimeError: If indexing fails
        """
        try:
            # Generate document_id if not provided, ensuring uniqueness
            if not document_id or document_id is None:
                document_id = await self._generate_unique_document_id(title)
                logger.info(
                    "Generated document_id from title",
                    document_id=document_id,
                    title=title
                )
            
            # Build base metadata
            full_metadata = metadata or {}
            full_metadata.update({
                "document_id": document_id,  # Now always has a value
                "title": title,
                "doc_type": doc_type.value,
                "industry": industry
            })
            
            # Automatically extract obligations and clauses if enabled
            if self.enable_auto_extraction and doc_type in [B2BDocumentType.LEASE, B2BDocumentType.CONTRACT]:
                try:
                    full_metadata = await self._extract_metadata_automatically(
                        text=text,
                        doc_type=doc_type,
                        metadata=full_metadata
                    )
                except Exception as e:
                    logger.warning(
                        "Auto-extraction failed, continuing with manual metadata",
                        error=str(e)
                    )
                    # Continue with manual metadata only
            
            # Chunk document
            documents = self._chunk_document(text, document_id, full_metadata)
            
            # Store in Milvus
            self.vector_store.add_documents(documents)
            
            logger.info(
                "Text indexed successfully",
                document_id=document_id,
                chunk_count=len(documents)
            )
            
            return IndexingResult(
                document_id=document_id,
                title=title,
                chunk_count=len(documents),
                format=DocumentFormat.TXT,  # Text content is treated as TXT
                file_size=len(text.encode('utf-8'))  # Approximate size in bytes
            )
            
        except Exception as e:
            logger.error(
                "Failed to index text",
                document_id=document_id,
                title=title,
                error=str(e),
                exc_info=True
            )
            raise RuntimeError(f"Failed to index text: {str(e)}") from e

    async def index_batch(
        self,
        files: List[Dict[str, Any]],
        industry: str = "generic",
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        Index multiple files in batch.
        
        Args:
            files: List of file info dicts, each containing:
                - file_path: Path to the file
                - document_id: Optional document ID
                - title: Document title
                - doc_type: Document type (B2BDocumentType enum or string)
                - metadata: Optional metadata dict
            industry: Industry classification for all documents
            chunk_size: Optional chunk size override
            chunk_overlap: Optional chunk overlap override
        
        Returns:
            Dict with:
                - total: Total number of files
                - success_count: Number of successfully indexed files
                - failed_count: Number of failed files
                - success_rate: Success rate (0.0 to 1.0)
                - successful: List of successful indexing results
                - failed: List of failed file info with errors
        """
        logger.info("Starting batch indexing", file_count=len(files), industry=industry)
        
        total = len(files)
        successful = []
        failed = []
        
        # Process files sequentially (can be parallelized later if needed)
        for i, file_info in enumerate(files):
            file_path = file_info.get("file_path")
            document_id = file_info.get("document_id")
            title = file_info.get("title", "Untitled Document")
            doc_type = file_info.get("doc_type")
            metadata = file_info.get("metadata", {})
            
            # Validate file_path
            if not file_path:
                failed.append({
                    "file_info": file_info,
                    "error": "file_path is required"
                })
                logger.warning("Skipping file: missing file_path", index=i)
                continue
            
            # Convert doc_type string to enum if needed
            if isinstance(doc_type, str):
                try:
                    doc_type_enum = B2BDocumentType(doc_type)
                except ValueError:
                    logger.warning(f"Unknown doc_type '{doc_type}', defaulting to LEGAL_DOCUMENT", file_path=file_path)
                    doc_type_enum = B2BDocumentType.LEGAL_DOCUMENT
            elif doc_type is None:
                doc_type_enum = B2BDocumentType.LEGAL_DOCUMENT
            elif isinstance(doc_type, B2BDocumentType):
                doc_type_enum = doc_type
            else:
                # Fallback for any other type
                doc_type_enum = B2BDocumentType.LEGAL_DOCUMENT
            
            # Merge metadata
            full_metadata = {**metadata}
            if not full_metadata.get("document_type"):
                full_metadata["document_type"] = doc_type_enum.value
            
            try:
                logger.info(
                    "Indexing file in batch",
                    index=i + 1,
                    total=total,
                    file_path=file_path,
                    title=title
                )
                
                # Index the file
                result = await self.index_file(
                    file_path=file_path,
                    document_id=document_id,
                    title=title,
                    doc_type=doc_type_enum,
                    industry=industry,
                    metadata=full_metadata if full_metadata else None
                )
                
                successful.append({
                    "document_id": result.document_id,
                    "title": result.title,
                    "chunk_count": result.chunk_count,
                    "format": result.format.value if hasattr(result.format, 'value') else str(result.format),
                    "file_path": file_path
                })
                
                logger.info(
                    "File indexed successfully in batch",
                    index=i + 1,
                    total=total,
                    document_id=result.document_id
                )
                
            except FileNotFoundError as e:
                error_msg = f"File not found: {file_path}"
                failed.append({
                    "file_path": file_path,
                    "title": title,
                    "error": error_msg
                })
                logger.error("File not found in batch", file_path=file_path, error=str(e))
                
            except ValueError as e:
                error_msg = f"Invalid file format or parameter: {str(e)}"
                failed.append({
                    "file_path": file_path,
                    "title": title,
                    "error": error_msg
                })
                logger.error("Invalid file in batch", file_path=file_path, error=str(e))
                
            except Exception as e:
                error_msg = f"Failed to index file: {str(e)}"
                failed.append({
                    "file_path": file_path,
                    "title": title,
                    "error": error_msg
                })
                logger.error(
                    "Failed to index file in batch",
                    file_path=file_path,
                    error=str(e),
                    exc_info=True
                )
        
        success_count = len(successful)
        failed_count = len(failed)
        success_rate = success_count / total if total > 0 else 0.0
        
        logger.info(
            "Batch indexing completed",
            total=total,
            success_count=success_count,
            failed_count=failed_count,
            success_rate=success_rate
        )
        
        return {
            "total": total,
            "success_count": success_count,
            "failed_count": failed_count,
            "success_rate": success_rate,
            "successful": successful,
            "failed": failed
        }

    async def search(
        self,
        query: str,
        doc_types: Optional[List[B2BDocumentType]] = None,
        industry: Optional[str] = None,
        top_k: int = 5,
        metadata_filters: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """
        Search indexed documents using semantic similarity.
        
        Performs vector similarity search to find the most relevant document chunks
        matching the query. Supports filtering by document type, industry, and metadata.
        
        Args:
            query: Search query text
            doc_types: Optional list of document types to filter by
            industry: Optional industry filter
            top_k: Number of results to return (default: 5)
            metadata_filters: Optional metadata filters (e.g., {"lease_id": "LEASE-001"})
        
        Returns:
            List of dictionaries containing:
            - content: Chunk text content
            - title: Document title
            - doc_type: Document type
            - document_id: Document identifier
            - chunk_index: Chunk index within document
            - score: Similarity score (0-1, higher = more relevant)
            - metadata: Full chunk metadata
        """
        try:
            # Build filter for document types if specified
            search_filter = None
            if doc_types:
                doc_type_values = [dt.value for dt in doc_types]
                # Milvus filter expression: doc_type in ["lease", "contract", ...]
                if len(doc_type_values) == 1:
                    search_filter = {"doc_type": doc_type_values[0]}
                else:
                    # For multiple types, we'll filter in post-processing
                    # (Milvus filter expressions can be complex)
                    search_filter = None
            
            # Add industry filter if specified
            if industry:
                if search_filter:
                    search_filter["industry"] = industry
                else:
                    search_filter = {"industry": industry}
            
            # Merge with additional metadata filters
            if metadata_filters:
                if search_filter:
                    search_filter.update(metadata_filters)
                else:
                    search_filter = metadata_filters
            
            # Perform similarity search
            documents = self.vector_store.similarity_search(
                query=query,
                k=top_k * 2,  # Get more results to filter if needed
                filters=search_filter
            )
            
            if not documents:
                logger.info("No documents found", query=query[:100])
                return []
            
            # Convert Document objects to dict format
            results = []
            for doc in documents:
                metadata = doc.metadata or {}
                
                # Debug logging to diagnose metadata issues
                logger.debug(
                    "Search result metadata",
                    metadata_keys=list(metadata.keys()),
                    has_document_id="document_id" in metadata,
                    document_id=metadata.get("document_id"),
                    doc_type=metadata.get("doc_type"),
                    title=metadata.get("title"),
                    content_preview=doc.page_content[:100] if doc.page_content else None
                )
                
                # Filter by doc_types if specified (post-processing)
                if doc_types:
                    doc_type_value = metadata.get("doc_type")
                    if doc_type_value not in [dt.value for dt in doc_types]:
                        continue
                
                # Filter by industry if specified (post-processing)
                if industry:
                    metadata_industry = metadata.get("industry")
                    if metadata_industry != industry:
                        continue
                
                # Apply additional metadata filters if any
                if metadata_filters:
                    matches = all(
                        metadata.get(key) == value
                        for key, value in metadata_filters.items()
                    )
                    if not matches:
                        continue
                
                # Calculate similarity score from distance
                # Milvus returns L2 distance (lower = more similar)
                # Convert distance to similarity: similarity = 1 / (1 + distance)
                # This gives a value between 0-1 where 1 = identical, 0 = very different
                distance = metadata.get("score", float('inf'))  # distance from Milvus
                if distance == float('inf'):
                    similarity_score = 0.0
                else:
                    # Convert distance to similarity (0-1 range)
                    # distance=0 → similarity=1.0 (100%), distance=1 → similarity=0.5 (50%)
                    similarity_score = 1.0 / (1.0 + distance)
                
                result = {
                    "content": doc.page_content,
                    "title": metadata.get("title", "Unknown"),
                    "doc_type": metadata.get("doc_type", "unknown"),
                    "document_id": metadata.get("document_id", "unknown"),
                    "chunk_index": metadata.get("chunk_index", 0),
                    "total_chunks": metadata.get("total_chunks", 0),
                    "score": similarity_score,  # Now a similarity score (0-1)
                    "metadata": metadata
                }
                results.append(result)
            
            # Limit to top_k after filtering
            results = results[:top_k]
            
            logger.info(
                "Search completed",
                query=query[:100],
                results_count=len(results),
                top_k=top_k
            )
            
            return results
            
        except Exception as e:
            logger.error(
                "Search failed",
                query=query[:100],
                error=str(e),
                exc_info=True
            )
            return []

    async def get_statistics(self) -> Dict[str, Any]:
        """
        Get statistics about indexed documents.
        
        Returns:
            Dictionary containing:
            - collection_name: Name of the Milvus collection
            - num_entities: Total number of document chunks indexed
            - milvus_available: Whether Milvus is connected
            - using_fallback: Whether using in-memory fallback
            - connection_state: Connection state
            - healthy: Overall health status
            - embeddings_available: Whether embeddings are available
        """
        try:
            # Get basic stats from vector store
            stats = self.vector_store.stats()
            
            # Add health check info
            health = self.vector_store.health_check()
            
            result = {
                "collection_name": self.collection_name,
                "num_entities": stats.get("num_entities", 0),
                "milvus_available": stats.get("milvus_available", False),
                "using_fallback": stats.get("using_fallback", False),
                "connection_state": stats.get("connection_state", "unknown"),
                "healthy": health.get("healthy", False),
                "embeddings_available": health.get("embeddings_available", False),
            }
            
            logger.info("Statistics retrieved", collection=self.collection_name, num_entities=result["num_entities"])
            return result
            
        except Exception as e:
            logger.error("Failed to get statistics", error=str(e), exc_info=True)
            return {
                "error": str(e),
                "collection_name": self.collection_name
            }

    async def delete_document(self, document_id: str) -> bool:
        """
        Delete a document and all its chunks from the vector store.
        
        Args:
            document_id: Unique document identifier
            
        Returns:
            True if deletion was successful, False otherwise
        """
        logger.info("Starting delete operation", document_id=document_id)
        try:
            # Check if chunks exist with this document_id
            test_docs = self.vector_store.list_documents(
                filters={"document_id": document_id},
                limit=10  # Get more to see what we're working with
            )
            
            logger.info(
                "Checked for chunks with document_id",
                document_id=document_id,
                chunks_found=len(test_docs)
            )
            
            if not test_docs:
                # Log what document_ids actually exist for debugging
                all_docs_sample = self.vector_store.list_documents(limit=20)
                sample_ids = set()
                for doc in all_docs_sample:
                    chunk_metadata = doc.get("metadata", {})
                    if isinstance(chunk_metadata, str):
                        try:
                            import json
                            chunk_metadata = json.loads(chunk_metadata)
                        except:
                            chunk_metadata = {}
                    doc_id = chunk_metadata.get("document_id")
                    if doc_id:
                        sample_ids.add(doc_id)
                
                logger.warning(
                    "Cannot delete: no chunks found with this document_id",
                    document_id=document_id,
                    sample_document_ids=list(sample_ids)[:10]  # Show first 10 for debugging
                )
                return False
            
            # Delete by document_id
            logger.info(
                "Deleting document by document_id",
                document_id=document_id,
                chunks_found=len(test_docs)
            )
            result = await self.vector_store.adelete_by_filter(
                filters={"document_id": document_id}
            )
            logger.info("Delete operation completed", document_id=document_id, result=result)
            
            # Milvus delete operations are asynchronous - wait a bit for flush
            import asyncio
            await asyncio.sleep(1.5)  # Increased wait time for Milvus to flush
            
            # Verify deletion
            verify_docs = self.vector_store.list_documents(
                filters={"document_id": document_id},
                limit=1
            )
            
            if verify_docs:
                logger.warning(
                    "Delete may have failed: chunks still exist after deletion",
                    document_id=document_id,
                    remaining_chunks=len(verify_docs)
                )
                return False
            
            logger.info("Document deleted successfully", document_id=document_id)
            return True
            
        except Exception as e:
            logger.error(
                "Failed to delete document",
                document_id=document_id,
                error=str(e),
                exc_info=True
            )
            return False

    async def clear_all_documents(self, confirm: bool = False) -> Dict[str, Any]:
        """
        Clear all documents from the vector store.
        
        WARNING: This will delete ALL indexed documents and cannot be undone!
        
        Args:
            confirm: Must be True to actually perform the deletion
            
        Returns:
            Dict with deletion statistics
        """
        if not confirm:
            return {
                "success": False,
                "message": "Deletion not confirmed. Set confirm=True to proceed.",
                "deleted": 0
            }
        
        logger.warning("Clearing all documents from vector store")
        
        try:
            # Get count before deletion
            stats_before = await self.get_statistics()
            total_before = stats_before.get("num_entities", 0)
            
            # Delete all documents by using a filter that matches everything
            # Since Milvus doesn't support delete all directly, we'll delete in batches
            # by getting all chunks and deleting them
            
            deleted_count = 0
            batch_size = 1000
            
            while True:
                # Get a batch of documents
                chunks = self.vector_store.list_documents(limit=batch_size)
                
                if not chunks:
                    break
                
                # Extract IDs
                chunk_ids = [chunk.get("id") for chunk in chunks if chunk.get("id")]
                
                if chunk_ids:
                    # Delete by IDs
                    result = await self.vector_store.adelete(chunk_ids)
                    deleted_count += len(chunk_ids)
                    logger.info("Deleted batch", count=len(chunk_ids), total_deleted=deleted_count)
                
                # If we got fewer than batch_size, we're done
                if len(chunks) < batch_size:
                    break
            
            # Wait for deletions to complete
            import asyncio
            await asyncio.sleep(2.0)
            
            # Verify deletion
            stats_after = await self.get_statistics()
            total_after = stats_after.get("num_entities", 0)
            
            return {
                "success": True,
                "message": f"Cleared {deleted_count} document chunks",
                "deleted": deleted_count,
                "total_before": total_before,
                "total_after": total_after
            }
            
        except Exception as e:
            logger.error("Failed to clear documents", error=str(e), exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "deleted": 0
            }

    async def list_indexed_documents(
        self,
        doc_type: Optional[str] = None,
        industry: Optional[str] = None,
        limit: int = 100,
        offset: int = 0
    ) -> Dict[str, Any]:
        """
        List indexed documents with optional filtering.
        
        Args:
            doc_type: Optional document type filter
            industry: Optional industry filter
            limit: Maximum number of documents to return
            offset: Number of documents to skip
            
        Returns:
            Dictionary containing:
            - documents: List of document summaries
            - total: Total count (if available)
        """
        try:
            # Build filters
            filters = {}
            if doc_type:
                filters["doc_type"] = doc_type
            if industry:
                filters["industry"] = industry
            
            # Get documents from vector store
            all_docs = self.vector_store.list_documents(
                filters=filters if filters else None,
                limit=limit * 10,  # Get more to deduplicate
                offset=offset
            )
            
            logger.debug(
                "Retrieved documents from vector store",
                count=len(all_docs),
                sample_doc=all_docs[0] if all_docs else None
            )
            
            # Deduplicate by document_id and aggregate chunk info
            document_map = {}
            skipped_no_id = 0
            for i, doc in enumerate(all_docs):
                # Parse metadata - handle both dict and JSON string
                raw_metadata = doc.get("metadata", {})
                if isinstance(raw_metadata, str):
                    try:
                        import json
                        metadata = json.loads(raw_metadata)
                    except (json.JSONDecodeError, TypeError):
                        logger.warning("Failed to parse metadata JSON", metadata=raw_metadata[:100] if raw_metadata else None)
                        metadata = {}
                else:
                    metadata = raw_metadata or {}
                
                # Log first few documents to debug metadata structure (use INFO so it shows up)
                if i < 3:
                    logger.info(
                        "Sample document metadata",
                        index=i,
                        metadata_keys=list(metadata.keys()) if isinstance(metadata, dict) else type(metadata),
                        metadata_preview=str(metadata)[:200] if metadata else None,
                        doc_keys=list(doc.keys()),
                        full_doc=str(doc)[:500]
                    )
                
                doc_id = metadata.get("document_id")
                
                # Only use stored document_id - don't generate new ones
                # All new documents should have document_id stored in chunks
                if not doc_id or doc_id is None:
                    # Try alternative field names as fallback
                    doc_id = metadata.get("doc_id") or metadata.get("id") or metadata.get("source")
                    
                    # If still no ID, skip this chunk (shouldn't happen for new documents)
                    if not doc_id or doc_id is None:
                        skipped_no_id += 1
                        if skipped_no_id <= 3:  # Only log first few
                            logger.warning(
                                "Skipping chunk without document_id (should not happen for new documents)",
                                metadata_keys=list(metadata.keys()) if isinstance(metadata, dict) else type(metadata),
                                metadata_preview=str(metadata)[:200] if metadata else None
                            )
                        continue
                
                if doc_id not in document_map:
                    document_map[doc_id] = {
                        "document_id": doc_id,
                        "title": metadata.get("title", "Unknown"),
                        "doc_type": metadata.get("doc_type", "unknown"),
                        "industry": metadata.get("industry", ""),
                        "chunk_count": 0
                    }
                
                document_map[doc_id]["chunk_count"] += 1
            
            # Convert to list and limit
            documents = list(document_map.values())[:limit]
            
            logger.info(
                "Listed documents",
                count=len(documents),
                total_chunks=len(all_docs),
                skipped_no_id=skipped_no_id,
                filters=filters
            )
            
            return {
                "documents": documents,
                "total": len(documents)
            }
            
        except Exception as e:
            logger.error(
                "Failed to list documents",
                error=str(e),
                exc_info=True
            )
            return {
                "documents": [],
                "total": 0
            }

    async def get_indexed_document(self, document_id: str) -> Optional[Dict[str, Any]]:
        """
        Get details about a specific indexed document.
        
        Args:
            document_id: Unique document identifier (may be generated from title)
            
        Returns:
            Dictionary with document details or None if not found
        """
        try:
            # Check if this is a generated document_id (starts with "doc_")
            # Generated IDs are created from titles when document_id is None in metadata
            if document_id.startswith("doc_"):
                # This is a generated ID - need to find the title first
                # List all documents to find the one with matching generated ID
                all_docs_list = await self.list_indexed_documents(limit=1000)
                matching_doc = None
                for doc in all_docs_list.get("documents", []):
                    if doc.get("document_id") == document_id:
                        matching_doc = doc
                        break
                
                if not matching_doc:
                    logger.warning(
                        "Cannot get document: document not found in list",
                        document_id=document_id
                    )
                    return None
                
                title = matching_doc.get("title")
                if not title:
                    logger.warning(
                        "Cannot get document: document has no title",
                        document_id=document_id
                    )
                    return None
                
                # Get all chunks with this title (since actual metadata has document_id: None)
                all_docs = self.vector_store.list_documents(
                    filters={"title": title},
                    limit=1000  # Get all chunks
                )
            else:
                # Normal document_id - get by document_id
                all_docs = self.vector_store.list_documents(
                    filters={"document_id": document_id},
                    limit=1000  # Get all chunks
                )
            
            if not all_docs:
                return None
            
            # Aggregate document info from chunks
            chunks = []
            metadata = {}
            
            for doc in all_docs:
                chunk_metadata = doc.get("metadata", {})
                if not metadata:
                    # Use first chunk's metadata as base
                    metadata = {
                        "document_id": chunk_metadata.get("document_id"),
                        "title": chunk_metadata.get("title", "Unknown"),
                        "doc_type": chunk_metadata.get("doc_type", "unknown"),
                        "industry": chunk_metadata.get("industry", ""),
                        **{k: v for k, v in chunk_metadata.items() 
                           if k not in ["chunk_index", "total_chunks", "document_id", "title", "doc_type", "industry"]}
                    }
                
                chunks.append({
                    "chunk_index": chunk_metadata.get("chunk_index", 0),
                    "content": doc.get("content", "")[:200],  # Preview
                    "metadata": chunk_metadata
                })
            
            # Sort chunks by index
            chunks.sort(key=lambda x: x["chunk_index"])
            
            # If this was a generated ID, use it in the result instead of None
            if document_id.startswith("doc_"):
                metadata["document_id"] = document_id
            
            result = {
                **metadata,
                "chunk_count": len(chunks),
                "chunks": chunks
            }
            
            logger.info("Retrieved document", document_id=document_id, chunk_count=len(chunks))
            return result
            
        except Exception as e:
            logger.error(
                "Failed to get document",
                document_id=document_id,
                error=str(e),
                exc_info=True
            )
            return None

    async def get_document_versions(self, document_id: str) -> Dict[str, Any]:
        """
        Get all versions of a document.
        
        Finds all documents that share the same contract_id, lease_id, or regulation_id
        as the given document_id, and returns them sorted by version/date.
        
        Args:
            document_id: Document ID to find versions for
            
        Returns:
            Dictionary containing:
            - base_document_id: The base identifier (contract_id, lease_id, etc.)
            - versions: List of version documents sorted by version/date
            - total_versions: Total number of versions found
        """
        try:
            # First, get the base document to find its contract_id/lease_id/regulation_id
            base_doc = await self.get_indexed_document(document_id)
            if not base_doc:
                logger.warning("Base document not found for versions", document_id=document_id)
                return {
                    "base_document_id": document_id,
                    "versions": [],
                    "total_versions": 0
                }
            
            # Determine the base identifier (contract_id, lease_id, or regulation_id)
            base_id = None
            base_id_type = None
            
            if base_doc.get("contract_id"):
                base_id = base_doc["contract_id"]
                base_id_type = "contract_id"
            elif base_doc.get("lease_id"):
                base_id = base_doc["lease_id"]
                base_id_type = "lease_id"
            elif base_doc.get("regulation_id"):
                base_id = base_doc["regulation_id"]
                base_id_type = "regulation_id"
            else:
                # If no base_id found, return just this document
                logger.info("No base_id found, returning single document", document_id=document_id)
                return {
                    "base_document_id": document_id,
                    "versions": [{
                        "document_id": document_id,
                        "title": base_doc.get("title", "Unknown"),
                        "version": base_doc.get("version", "1.0"),
                        "version_date": base_doc.get("version_date"),
                        "previous_version_id": base_doc.get("previous_version_id"),
                        "chunk_count": base_doc.get("chunk_count", 0)
                    }],
                    "total_versions": 1
                }
            
            # Query all documents and filter in Python (Milvus JSON filtering is unreliable)
            # Get a large batch of documents to filter
            all_docs_raw = self.vector_store.list_documents(
                filters=None,  # Don't filter at Milvus level - JSON filtering doesn't work reliably
                limit=10000  # Get a large batch
            )
            
            # Filter in Python by base_id
            all_docs = []
            for doc in all_docs_raw:
                chunk_metadata = doc.get("metadata", {})
                if isinstance(chunk_metadata, str):
                    try:
                        import json
                        chunk_metadata = json.loads(chunk_metadata)
                    except:
                        chunk_metadata = {}
                
                # Check if this chunk's metadata matches our base_id
                if chunk_metadata.get(base_id_type) == base_id:
                    all_docs.append(doc)
            
            # Deduplicate by document_id and extract version info
            version_map = {}
            for doc in all_docs:
                chunk_metadata = doc.get("metadata", {})
                if isinstance(chunk_metadata, str):
                    try:
                        import json
                        chunk_metadata = json.loads(chunk_metadata)
                    except:
                        chunk_metadata = {}
                
                doc_id = chunk_metadata.get("document_id")
                if not doc_id:
                    continue
                
                if doc_id not in version_map:
                    version_map[doc_id] = {
                        "document_id": doc_id,
                        "title": chunk_metadata.get("title", "Unknown"),
                        "version": chunk_metadata.get("version", "1.0"),
                        "version_date": chunk_metadata.get("version_date"),
                        "previous_version_id": chunk_metadata.get("previous_version_id"),
                        "chunk_count": 0
                    }
                
                version_map[doc_id]["chunk_count"] += 1
            
            # Sort versions by version_date or version number
            versions = list(version_map.values())
            try:
                # Try to sort by version_date first, then by version string
                versions.sort(key=lambda v: (
                    v.get("version_date") or "",
                    v.get("version") or "0"
                ))
            except:
                # Fallback: sort by version string
                versions.sort(key=lambda v: v.get("version") or "0")
            
            logger.info(
                "Retrieved document versions",
                document_id=document_id,
                base_id=base_id,
                base_id_type=base_id_type,
                version_count=len(versions)
            )
            
            return {
                "base_document_id": base_id,
                "base_id_type": base_id_type,
                "versions": versions,
                "total_versions": len(versions)
            }
            
        except Exception as e:
            logger.error(
                "Failed to get document versions",
                document_id=document_id,
                error=str(e),
                exc_info=True
            )
            return {
                "base_document_id": document_id,
                "versions": [],
                "total_versions": 0,
                "error": str(e)
            }

    async def get_document_obligations(self, document_id: str) -> Dict[str, Any]:
        """
        Get obligations for a specific document.
        
        Extracts obligations from document metadata that were either:
        - Manually provided during indexing
        - Auto-extracted during indexing
        
        Args:
            document_id: Document ID to get obligations for
            
        Returns:
            Dictionary containing:
            - document_id: The document ID
            - title: Document title
            - obligations: List of obligations
            - total_obligations: Total number of obligations
        """
        try:
            # Get the document
            doc = await self.get_indexed_document(document_id)
            if not doc:
                logger.warning("Document not found for obligations", document_id=document_id)
                return {
                    "document_id": document_id,
                    "title": None,
                    "obligations": [],
                    "total_obligations": 0
                }
            
            # Extract obligations from metadata
            obligations = doc.get("obligations", [])
            
            # If obligations is a string, try to parse it
            if isinstance(obligations, str):
                try:
                    import json
                    obligations = json.loads(obligations)
                except:
                    obligations = []
            
            # Ensure obligations is a list
            if not isinstance(obligations, list):
                obligations = []
            
            logger.info(
                "Retrieved document obligations",
                document_id=document_id,
                obligation_count=len(obligations)
            )
            
            return {
                "document_id": document_id,
                "title": doc.get("title", "Unknown"),
                "obligations": obligations,
                "total_obligations": len(obligations)
            }
            
        except Exception as e:
            logger.error(
                "Failed to get document obligations",
                document_id=document_id,
                error=str(e),
                exc_info=True
            )
            return {
                "document_id": document_id,
                "title": None,
                "obligations": [],
                "total_obligations": 0,
                "error": str(e)
            }

# Singleton instance
_document_indexing_service: Optional[DocumentIndexingService] = None

async def get_document_indexing_service(
    collection_name: str = "language_intelligence_documents",
    **kwargs
) -> DocumentIndexingService:
    """
    Get singleton DocumentIndexingService instance.
    
    Args:
        collection_name: Name of the Milvus collection
        **kwargs: Additional arguments passed to DocumentIndexingService.__init__
    
    Returns:
        DocumentIndexingService instance
    """
    global _document_indexing_service
    if _document_indexing_service is None:
        _document_indexing_service = DocumentIndexingService(
            collection_name=collection_name,
            **kwargs
        )
    return _document_indexing_service